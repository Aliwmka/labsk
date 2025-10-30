from docx import Document
import pandas as pd
import tkinter as tk
from tkinter import filedialog, messagebox
from sql_requests import (
    INSERT_TEACHER_SQL, INSERT_SUBJECT_SQL, INSERT_WORKLOAD_SQL,
    EDIT_TEACHER_SQL, EDIT_SUBJECT_SQL, EDIT_WORKLOAD_SQL
)

def export_to_word(instance):
    # Получаем данные из таблицы
    rows = [(instance.results_table.item(item)["values"]) for item in instance.results_table.get_children()]
    
    # Проверяем, есть ли данные для экспорта
    if not rows:
        messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
        return
    
    # Создаем документ Word
    doc = Document()
    doc.add_heading('Отчет по распределению учебной нагрузки', 0)
    
    # Добавляем таблицу
    table = doc.add_table(rows=1, cols=len(instance.results_table["columns"]))
    hdr_cells = table.rows[0].cells
    for i, column in enumerate(instance.results_table["columns"]):
        hdr_cells[i].text = column
    
    for row in rows:
        row_cells = table.add_row().cells
        for i, value in enumerate(row):
            row_cells[i].text = str(value)
    
    # Сохраняем документ
    teacher_option = instance.teacher_combobox.get() or "Все"
    subject_option = instance.subject_combobox.get() or "Все"
    filename = f"Отчет_нагрузка_{teacher_option}_{subject_option}.docx".replace(" ", "_")
    doc.save(filename)
    messagebox.showinfo("Успех", f"Отчет сохранен в {filename}")

def export_to_excel(instance):
    # Получаем данные из таблицы
    rows = [(instance.results_table.item(item)["values"]) for item in instance.results_table.get_children()]
    columns = instance.results_table["columns"]
    
    # Проверяем, есть ли данные для экспорта
    if not rows:
        messagebox.showwarning("Предупреждение", "Нет данных для экспорта")
        return
    
    # Создаем DataFrame и сохраняем в Excel
    df = pd.DataFrame(rows, columns=columns)
    teacher_option = instance.teacher_combobox.get() or "Все"
    subject_option = instance.subject_combobox.get() or "Все"
    filename = f"Отчет_нагрузка_{teacher_option}_{subject_option}.xlsx".replace(" ", "_")
    df.to_excel(filename, index=False)
    messagebox.showinfo("Успех", f"Отчет сохранен в {filename}")

def import_from_excel(instance, table_type):
    """Импорт данных из Excel файла"""
    file_path = filedialog.askopenfilename(
        title="Выберите файл Excel для импорта",
        filetypes=[("Excel files", "*.xlsx *.xls")]
    )
    
    if not file_path:
        return
    
    try:
        df = pd.read_excel(file_path)
        
        if table_type == "teachers":
            import_teachers_from_df(instance, df)
        elif table_type == "subjects":
            import_subjects_from_df(instance, df)
        elif table_type == "workload":
            import_workload_from_df(instance, df)
        
        messagebox.showinfo("Успех", "Данные успешно импортированы!")
        
    except Exception as e:
        messagebox.showerror("Ошибка", f"Ошибка при импорте данных: {str(e)}")

def import_teachers_from_df(instance, df):
    """Импорт преподавателей из DataFrame"""
    required_columns = ['Фамилия', 'Имя', 'Отчество', 'Ученая степень', 'Должность', 'Стаж']
    
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"Отсутствует обязательная колонка: {col}")
    
    imported_count = 0
    updated_count = 0
    
    for _, row in df.iterrows():
        data = (
            str(row['Фамилия']), str(row['Имя']), str(row['Отчество'] or ''),
            str(row['Ученая степень'] or ''), str(row['Должность'] or ''), 
            int(row['Стаж']) if pd.notna(row['Стаж']) else 0
        )
        
        # Проверяем, существует ли уже такой преподаватель
        cur = instance.db_connection.cursor()
        cur.execute(
            "SELECT teacher_id FROM teachers WHERE last_name = %s AND first_name = %s AND middle_name = %s",
            (data[0], data[1], data[2])
        )
        existing = cur.fetchone()
        
        if existing:
            # Обновляем существующую запись
            cur.execute(EDIT_TEACHER_SQL, (*data, existing[0]))
            updated_count += 1
        else:
            # Добавляем новую запись
            cur.execute(INSERT_TEACHER_SQL, data)
            imported_count += 1
        
        instance.db_connection.commit()
        cur.close()
    
    messagebox.showinfo("Импорт преподавателей", 
                       f"Импортировано: {imported_count} записей\nОбновлено: {updated_count} записей")

def import_subjects_from_df(instance, df):
    """Импорт предметов из DataFrame"""
    required_columns = ['Название предмета', 'Количество часов']
    
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"Отсутствует обязательная колонка: {col}")
    
    imported_count = 0
    updated_count = 0
    
    for _, row in df.iterrows():
        data = (
            str(row['Название предмета']), 
            int(row['Количество часов']) if pd.notna(row['Количество часов']) else 0
        )
        
        # Проверяем, существует ли уже такой предмет
        cur = instance.db_connection.cursor()
        cur.execute(
            "SELECT subject_id FROM subjects WHERE subject_name = %s",
            (data[0],)
        )
        existing = cur.fetchone()
        
        if existing:
            # Обновляем существующую запись
            cur.execute(EDIT_SUBJECT_SQL, (*data, existing[0]))
            updated_count += 1
        else:
            # Добавляем новую запись
            cur.execute(INSERT_SUBJECT_SQL, data)
            imported_count += 1
        
        instance.db_connection.commit()
        cur.close()
    
    messagebox.showinfo("Импорт предметов", 
                       f"Импортировано: {imported_count} записей\nОбновлено: {updated_count} записей")

def import_workload_from_df(instance, df):
    """Импорт нагрузки из DataFrame"""
    required_columns = ['Преподаватель', 'Предмет', 'Группа']
    
    for col in required_columns:
        if col not in df.columns:
            raise ValueError(f"Отсутствует обязательная колонка: {col}")
    
    imported_count = 0
    updated_count = 0
    skipped_count = 0
    
    for _, row in df.iterrows():
        teacher_name = str(row['Преподаватель'])
        subject_name = str(row['Предмет'])
        group_number = str(row['Группа'])
        
        # Получаем ID преподавателя
        cur = instance.db_connection.cursor()
        cur.execute(
            "SELECT teacher_id FROM teachers WHERE CONCAT(last_name, ' ', first_name, ' ', middle_name) = %s",
            (teacher_name,)
        )
        teacher_result = cur.fetchone()
        
        if not teacher_result:
            skipped_count += 1
            continue
        
        teacher_id = teacher_result[0]
        
        # Получаем ID предмета
        cur.execute(
            "SELECT subject_id FROM subjects WHERE subject_name = %s",
            (subject_name,)
        )
        subject_result = cur.fetchone()
        
        if not subject_result:
            skipped_count += 1
            continue
        
        subject_id = subject_result[0]
        
        # Проверяем, существует ли уже такая нагрузка
        cur.execute(
            "SELECT workload_id FROM workload WHERE teacher_id = %s AND subject_id = %s AND group_number = %s",
            (teacher_id, subject_id, group_number)
        )
        existing = cur.fetchone()
        
        if existing:
            # Обновляем существующую запись
            cur.execute(EDIT_WORKLOAD_SQL, (teacher_id, subject_id, group_number, existing[0]))
            updated_count += 1
        else:
            # Добавляем новую запись
            cur.execute(INSERT_WORKLOAD_SQL, (teacher_id, subject_id, group_number))
            imported_count += 1
        
        instance.db_connection.commit()
        cur.close()
    
    messagebox.showinfo("Импорт нагрузки", 
                       f"Импортировано: {imported_count} записей\n"
                       f"Обновлено: {updated_count} записей\n"
                       f"Пропущено (не найдены преподаватели/предметы): {skipped_count} записей")